"""Document reading and editing tools (Word, Excel, PDF).

v1.7 「Office 体系 2.0」: write_document / write_excel were v0.x 老工具, never wired into the design-token
system. On a 中文 Windows box that produced docx/xlsx with ZERO explicit font declarations, so Word/Excel
fell back to per-level defaults (拉丁 Calibri / 中文 宋体 混排, 标题各回落各的 —— 用户真机实测「字体都不对，
中文大小不一」). Root cause for Word: python-docx's `font.name` only writes the ascii/hAnsi rFonts slots;
the CJK glyph run is governed by the SEPARATE `w:eastAsia` rFonts attribute, which python-docx never
touches — so 中文 always falls back. Fix (this module): inject explicit fonts at the *styles.xml* level
(Normal / Title / Heading 1-3), writing all three rFonts links (ascii / hAnsi / **eastAsia**) to the
token font, plus a size ladder and heading colour, via a hand-written lxml `_set_style_font`. No run is
left 裸奔 — every paragraph inherits a fully-specified style. write_excel gets the same treatment (每格
落笔即带 token Font + 数字格式启发式), so 「不跑 beautify 也不难看」.
"""

import os
from ai_computer_control.server import mcp
from ai_computer_control.tools.safety import protected_path_reason
from ai_computer_control.tools import office_style as style_tokens


# v1.0 收官安全加固(对抗复核 CONFIRMED·minor):写族(write_document/write_excel/write_pdf)此前不接
# protected_path_reason,而删/移/拷族(filesystem.py)都接 —— 造成「能往受保护系统树种植/覆盖文件、却删不掉」
# 的护栏不对称。补齐:写前对目标路径过同一护栏,带 allow_protected 逃生阀,与 delete/move/copy 一致。
def _protected_write_guard(path: str, allow_protected: bool):
    reason = protected_path_reason(path)
    if reason and not allow_protected:
        return {"error": f"refused: destination {reason}. Pass allow_protected=true to override."}
    return None


@mcp.tool()
def read_document(path: str) -> dict:
    """Read the text content of a document file (Word .docx, Excel .xlsx, PDF .pdf).

    Args:
        path: Path to the document file.

    Returns:
        dict with 'content' (extracted text), 'type', 'pages'/'sheets' count.
    """
    ext = os.path.splitext(path)[1].lower()

    try:
        if ext == ".docx":
            return _read_docx(path)
        elif ext == ".xlsx":
            return _read_xlsx(path)
        elif ext == ".pdf":
            return _read_pdf(path)
        else:
            return {"error": f"Unsupported format: {ext}. Supported: .docx, .xlsx, .pdf"}
    except Exception as e:
        return {"error": str(e)}


def _read_docx(path: str) -> dict:
    from docx import Document

    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs]

    # Also read tables
    tables_text = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            rows.append(" | ".join(cells))
        tables_text.append("\n".join(rows))

    content = "\n".join(paragraphs)
    if tables_text:
        content += "\n\n--- Tables ---\n" + "\n\n".join(tables_text)

    return {
        "content": content,
        "type": "docx",
        "paragraphs": len(paragraphs),
        "tables": len(doc.tables),
    }


def _read_xlsx(path: str) -> dict:
    from openpyxl import load_workbook

    wb = load_workbook(path, read_only=True, data_only=True)
    sheets = {}

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            rows.append([str(cell) if cell is not None else "" for cell in row])
        sheets[sheet_name] = rows

    wb.close()

    content_parts = []
    for name, rows in sheets.items():
        content_parts.append(f"=== Sheet: {name} ===")
        for row in rows[:500]:
            content_parts.append(" | ".join(row))

    return {
        "content": "\n".join(content_parts),
        "type": "xlsx",
        "sheets": list(sheets.keys()),
        "sheet_count": len(sheets),
    }


def _read_pdf(path: str) -> dict:
    import pdfplumber

    text_parts = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text()
            if text:
                text_parts.append(f"--- Page {i + 1} ---\n{text}")

        page_count = len(pdf.pages)

    return {
        "content": "\n\n".join(text_parts),
        "type": "pdf",
        "pages": page_count,
    }


# =============================================================================================
# Word 字体纪律 helpers (v1.7 「Office 体系 2.0」)
#
# The生死线: python-docx sets ascii/hAnsi rFonts via Font.name but NEVER w:eastAsia, so on 中文
# Windows every CJK run falls back to the theme's minor-EA font (宋体). We hand-write the whole
# w:rPr/w:rFonts (ascii + hAnsi + eastAsia all = token font) at the STYLE level in styles.xml, plus
# the size / colour / bold, via lxml directly on `style.element.get_or_add_rPr()`. Applying this to
# Normal / Title / Heading 1-3 means every paragraph that uses a built-in style inherits a fully
# specified font — no run is ever 裸奔 (font-declaration-free) again.
# =============================================================================================

# Size ladder (pt) —用户拍板: Normal 11 / H3 14 / H2 16 / H1 20 / Title 28.
_WORD_SIZES = {"Normal": 11, "Heading 3": 14, "Heading 2": 16, "Heading 1": 20, "Title": 28}


def _set_style_font(style, name, size_pt, color_hex=None, bold=None):
    """Force a python-docx *style* onto an explicit font at the styles.xml (rPr) level via lxml.

    Writes, on the style's `<w:rPr>`:
      * `<w:rFonts w:ascii w:hAnsi w:eastAsia w:cs>` — ALL FOUR set to `name` so 拉丁 + 中文 use the
        same family (the eastAsia slot is the one python-docx never sets → the whole 中文大小不一 bug).
      * `<w:sz>` / `<w:szCs>` — size in half-points (size_pt * 2).
      * `<w:color>` — hex colour (no '#'), when color_hex is given.
      * `<w:b>` / `<w:bCs>` — bold on/off, when bold is not None.

    Idempotent: existing rFonts/sz/color/b children are removed and re-appended, so re-styling a
    style never stacks duplicate elements.
    """
    from docx.oxml.ns import qn

    rpr = style.element.get_or_add_rPr()

    def _clear(tag):
        for el in rpr.findall(qn(tag)):
            rpr.remove(el)

    # rFonts: ascii / hAnsi / eastAsia / cs all = name.
    _clear("w:rFonts")
    rfonts = rpr.makeelement(qn("w:rFonts"), {})
    for slot in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(slot), name)
    rpr.append(rfonts)

    # sz / szCs in half-points.
    _clear("w:sz")
    _clear("w:szCs")
    half = str(int(round(size_pt * 2)))
    for tag in ("w:sz", "w:szCs"):
        el = rpr.makeelement(qn(tag), {})
        el.set(qn("w:val"), half)
        rpr.append(el)

    if color_hex is not None:
        _clear("w:color")
        col = rpr.makeelement(qn("w:color"), {})
        col.set(qn("w:val"), str(color_hex).lstrip("#").upper())
        rpr.append(col)

    if bold is not None:
        _clear("w:b")
        _clear("w:bCs")
        for tag in ("w:b", "w:bCs"):
            el = rpr.makeelement(qn(tag), {})
            el.set(qn("w:val"), "1" if bold else "0")
            rpr.append(el)


def _apply_word_styles(doc, tokens):
    """Inject token fonts/sizes/colours into Normal / Title / Heading 1-3 (the 字体纪律 backbone).

    Normal: body font, 11pt, text colour, 行距 1.4. Title: title font, 28pt, deep title colour, bold.
    H1/H2/H3: heading font, 20/16/14pt, heading colour, bold. Heading paragraphs also get a coloured
    hairline bottom border (2.25pt) via _add_heading_rule at content time (border lives on the
    paragraph, not the style, so python-docx can't carry it in styles.xml reliably).
    """
    from docx.shared import Pt

    body_font = tokens["body_font"]
    title_font = tokens["title_font"]
    text_color = tokens["text_color"]
    title_color = tokens["word_title_color"]
    heading_color = tokens["word_heading_color"]

    styles = doc.styles

    # Normal — the parent of everything; also set line spacing 1.4 + a touch of space-after.
    normal = styles["Normal"]
    _set_style_font(normal, body_font, _WORD_SIZES["Normal"], color_hex=text_color, bold=False)
    pf = normal.paragraph_format
    pf.line_spacing = 1.4
    pf.space_after = Pt(6)

    # Title (level-0 heading via add_heading(..,0)).
    try:
        _set_style_font(styles["Title"], title_font, _WORD_SIZES["Title"],
                        color_hex=title_color, bold=True)
        styles["Title"].paragraph_format.space_after = Pt(10)
    except KeyError:
        pass

    for lvl, sz in (("Heading 1", _WORD_SIZES["Heading 1"]),
                    ("Heading 2", _WORD_SIZES["Heading 2"]),
                    ("Heading 3", _WORD_SIZES["Heading 3"])):
        try:
            st = styles[lvl]
            _set_style_font(st, title_font, sz, color_hex=heading_color, bold=True)
            hpf = st.paragraph_format
            # H1 前段距加大 so sections breathe; H2/H3 modest.
            hpf.space_before = Pt(18 if lvl == "Heading 1" else 12)
            hpf.space_after = Pt(4)
        except KeyError:
            pass

    # List Bullet / List Number: these inherit from Normal but their own rPr may not carry the
    # eastAsia link, so a bullet run can still fall back to 宋体. Inject the body font explicitly at
    # Normal 字号 so 无裸 run also holds for list items (the 字体纪律 must cover 项目符号).
    for lst in ("List Bullet", "List Number"):
        try:
            _set_style_font(styles[lst], body_font, _WORD_SIZES["Normal"],
                            color_hex=text_color, bold=False)
        except KeyError:
            pass


def _add_heading_rule(paragraph, color_hex, side="bottom", sz="18"):
    """Add a coloured hairline as a paragraph border (标题强调线).

    side='bottom' (default) — 标题下细横线 (business/vibrant v1.7 原样式; sz 18 = 2.25pt).
    side='left'   (v1.7.1 minimal) — 左侧青色细竖线 (段落左边框), 墨白极简的「装帧竖线」语言:
                  不用底纹/横幅, 仅一条竖 hairline 标记层级. sz 单位是 1/8 pt.
    Idempotent: an existing pBdr is replaced, never stacked."""
    from docx.oxml.ns import qn

    ppr = paragraph._p.get_or_add_pPr()
    for existing in ppr.findall(qn("w:pBdr")):
        ppr.remove(existing)
    pbdr = ppr.makeelement(qn("w:pBdr"), {})
    edge = pbdr.makeelement(qn("w:" + ("left" if side == "left" else "bottom")), {})
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), str(sz))      # eighths of a point → 18 = 2.25pt
    edge.set(qn("w:space"), "4")
    edge.set(qn("w:color"), str(color_hex).lstrip("#").upper())
    pbdr.append(edge)
    ppr.append(pbdr)


def _add_page_number_footer(doc, tokens):
    """Centre a 「第 X 页」 footer using a PAGE field so Word renders the live page number."""
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _run_text(text):
        r = p.add_run(text)
        r.font.name = tokens["body_font"]
        r.font.size = Pt(9)
        # eastAsia on the run too (footer runs are direct, not style-driven).
        rpr = r._element.get_or_add_rPr()
        rf = rpr.find(qn("w:rFonts"))
        if rf is None:
            rf = rpr.makeelement(qn("w:rFonts"), {})
            rpr.append(rf)
        rf.set(qn("w:eastAsia"), tokens["body_font"])
        return r

    def _field(instr):
        r = p.add_run()
        fld_begin = r._element.makeelement(qn("w:fldChar"), {})
        fld_begin.set(qn("w:fldCharType"), "begin")
        r._element.append(fld_begin)
        r2 = p.add_run()
        instr_el = r2._element.makeelement(qn("w:instrText"), {})
        instr_el.set(qn("xml:space"), "preserve")
        instr_el.text = instr
        r2._element.append(instr_el)
        r3 = p.add_run()
        fld_end = r3._element.makeelement(qn("w:fldChar"), {})
        fld_end.set(qn("w:fldCharType"), "end")
        r3._element.append(fld_end)

    _run_text("第 ")
    _field(" PAGE ")
    _run_text(" 页")


def _add_cover(doc, cover, tokens):
    """Render the cover page then a page break. v1.7.1 dispatches on the word_cover 版式选择器:

      * 'dark_block' (business/vibrant, v1.7 原样) — 深底满铺单元格题块 + 白大字 + 金/强调线;
        vibrant 额外在题块下沿加一条 0.3in 珊瑚粗条 (word_cover_bar token).
      * 'light_top'  (minimal, v1.7.1 new) — 白底装帧: 墨黑特大字置于页面上 1/3, 下方一条细青线,
        副题灰. 无任何色块 (高级文印/咨询装帧).

    cover: {title, subtitle?, date?, author?}.
    """
    if tokens.get("word_cover", "dark_block") == "light_top":
        _add_cover_light(doc, cover, tokens)
        return

    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.shared import Pt, RGBColor

    bg = tokens["word_cover_bg"]
    fg = tokens["word_cover_fg"]
    sub = tokens["word_cover_sub"]
    accent = tokens["accent"]           # 鎏金 on the dark cover block
    title_font = tokens["title_font"]
    body_font = tokens["body_font"]

    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    # shade the cell with the deep cover colour.
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.makeelement(qn("w:shd"), {})
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), bg.lstrip("#").upper())
    tcPr.append(shd)
    # remove default table borders (borderless block).
    tblPr = tbl._tbl.tblPr
    borders = tblPr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = borders.makeelement(qn("w:" + edge), {})
        b.set(qn("w:val"), "none")
        borders.append(b)
    tblPr.append(borders)

    def _styled_run(paragraph, text, size, color_hex, bold=False):
        r = paragraph.add_run(text)
        r.font.name = title_font if bold else body_font
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.color.rgb = RGBColor(*style_tokens.rgb_tuple(color_hex))
        rpr = r._element.get_or_add_rPr()
        rf = rpr.find(qn("w:rFonts"))
        if rf is None:
            rf = rpr.makeelement(qn("w:rFonts"), {})
            rpr.append(rf)
        rf.set(qn("w:eastAsia"), title_font if bold else body_font)
        return r

    # top spacer
    sp = cell.paragraphs[0]
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp.paragraph_format.space_before = Pt(48)

    # title
    tp = cell.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_before = Pt(6)
    _styled_run(tp, str(cover.get("title", "")), 32, fg, bold=True)

    # accent 鎏金 rule under the title (drawn on the dark block → 鎏金 shines here)
    _add_heading_rule(tp, accent)

    subtitle = cover.get("subtitle")
    if subtitle:
        subp = cell.add_paragraph()
        subp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subp.paragraph_format.space_before = Pt(10)
        _styled_run(subp, str(subtitle), 16, sub, bold=False)

    meta = []
    if cover.get("author"):
        meta.append(str(cover.get("author")))
    if cover.get("date"):
        meta.append(str(cover.get("date")))
    if meta:
        mp = cell.add_paragraph()
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mp.paragraph_format.space_before = Pt(24)
        mp.paragraph_format.space_after = Pt(48)
        _styled_run(mp, "　·　".join(meta), 12, sub, bold=False)
    else:
        # bottom spacer so the block has vertical breathing room even without meta
        bp = cell.add_paragraph()
        bp.paragraph_format.space_after = Pt(48)

    # v1.7.1 vibrant: 珊瑚粗条 (0.3in) 在题块下沿 — a second exact-height row shaded coral appended
    # to the same borderless table (adjacent tables merge in Word, so a row IS the clean way).
    bar_hex = tokens.get("word_cover_bar")
    if bar_hex:
        row2 = tbl.add_row()
        bar_cell = row2.cells[0]
        btcPr = bar_cell._tc.get_or_add_tcPr()
        bshd = btcPr.makeelement(qn("w:shd"), {})
        bshd.set(qn("w:val"), "clear")
        bshd.set(qn("w:color"), "auto")
        bshd.set(qn("w:fill"), str(bar_hex).lstrip("#").upper())
        btcPr.append(bshd)
        # exact row height 0.3in = 432 twips; squash the cell paragraph so it can't stretch the row.
        trPr = row2._tr.get_or_add_trPr()
        trH = trPr.makeelement(qn("w:trHeight"), {})
        trH.set(qn("w:val"), "432")
        trH.set(qn("w:hRule"), "exact")
        trPr.append(trH)
        bp2 = bar_cell.paragraphs[0]
        bp2.paragraph_format.space_before = Pt(0)
        bp2.paragraph_format.space_after = Pt(0)
        bp2.paragraph_format.line_spacing = 1.0
        br = bp2.add_run("")
        br.font.size = Pt(2)

    # page break so the body starts on a fresh page
    doc.add_page_break()


def _add_cover_light(doc, cover, tokens):
    """minimal「墨白极简」Word 封面 (v1.7.1) — 高级文印/咨询装帧:

    白底无色块; 标题墨黑特大字 (36pt) 置于页面上 1/3 (上方留白约 2in), 左对齐 (与 minimal PPT 封面的
    左对齐语言一致); 标题下一条细青线 (1.5pt, 段落下边框); 副题灰; 作者·日期灰小字。然后分页。
    """
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    ink = tokens["word_title_color"]        # 222222 墨黑
    accent = tokens["accent"]               # 3B7C8C 青
    sub = tokens["subtle_color"]            # 8C8C8C 灰 (白底上的次要字色)
    title_font = tokens["title_font"]
    body_font = tokens["body_font"]

    def _styled_run(paragraph, text, size, color_hex, bold=False):
        r = paragraph.add_run(text)
        r.font.name = title_font if bold else body_font
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.color.rgb = RGBColor(*style_tokens.rgb_tuple(color_hex))
        rpr = r._element.get_or_add_rPr()
        rf = rpr.find(qn("w:rFonts"))
        if rf is None:
            rf = rpr.makeelement(qn("w:rFonts"), {})
            rpr.append(rf)
        rf.set(qn("w:eastAsia"), title_font if bold else body_font)
        return r

    # push the title block down to the top-third of the page (~2in of air under the top margin).
    spacer = doc.add_paragraph("")
    spacer.paragraph_format.space_after = Pt(144)

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _styled_run(tp, str(cover.get("title", "")), 36, ink, bold=True)
    # 细青线 under the title — thinner than the business rule (12 = 1.5pt), same pBdr mechanism.
    _add_heading_rule(tp, accent, side="bottom", sz="12")

    subtitle = cover.get("subtitle")
    if subtitle:
        subp = doc.add_paragraph()
        subp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        subp.paragraph_format.space_before = Pt(14)
        _styled_run(subp, str(subtitle), 15, sub, bold=False)

    meta = []
    if cover.get("author"):
        meta.append(str(cover.get("author")))
    if cover.get("date"):
        meta.append(str(cover.get("date")))
    if meta:
        mp = doc.add_paragraph()
        mp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        mp.paragraph_format.space_before = Pt(30)
        _styled_run(mp, "　·　".join(meta), 11, sub, bold=False)

    doc.add_page_break()


def _add_content_table(doc, headers, rows, tokens):
    """Add a token-styled table (表头主色填充白字 + 斑马纹 + 细边框), matching the Excel/PPT observation."""
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    primary = tokens["header_fill"]
    header_fg = tokens["header_font_color"]
    zebra = tokens["zebra_fill"]
    border = tokens["border_color"]
    text_color = tokens["text_color"]
    body_font = tokens["body_font"]

    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.style = "Table Grid"  # gives us a real grid to recolour

    # recolour all borders to the token hairline
    tblPr = tbl._tbl.tblPr
    for existing in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(existing)
    borders = tblPr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = borders.makeelement(qn("w:" + edge), {})
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")            # 0.5pt
        b.set(qn("w:color"), border.lstrip("#").upper())
        borders.append(b)
    tblPr.append(borders)

    def _shade(cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        for ex in tcPr.findall(qn("w:shd")):
            tcPr.remove(ex)
        shd = tcPr.makeelement(qn("w:shd"), {})
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill_hex.lstrip("#").upper())
        tcPr.append(shd)

    def _cell_run(cell, text, color_hex, bold):
        cell.paragraphs[0].text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(str(text))
        r.font.name = body_font
        r.font.size = Pt(10.5)
        r.font.bold = bold
        r.font.color.rgb = RGBColor(*style_tokens.rgb_tuple(color_hex))
        rpr = r._element.get_or_add_rPr()
        rf = rpr.find(qn("w:rFonts"))
        if rf is None:
            rf = rpr.makeelement(qn("w:rFonts"), {})
            rpr.append(rf)
        rf.set(qn("w:eastAsia"), body_font)

    # header row
    for c, h in enumerate(headers):
        cell = tbl.cell(0, c)
        _shade(cell, primary)
        _cell_run(cell, h, header_fg, bold=True)

    # body rows with zebra
    for ri, row in enumerate(rows, start=1):
        striped = (ri % 2 == 0)
        for c in range(n_cols):
            cell = tbl.cell(ri, c)
            if striped:
                _shade(cell, zebra)
            val = row[c] if c < len(row) else ""
            _cell_run(cell, val, text_color, bold=False)


@mcp.tool()
def write_document(
    path: str,
    content: str,
    title: str | None = None,
    style: str = "business",
    cover: dict | None = None,
    page_numbers: bool = False,
    allow_protected: bool = False,
) -> dict:
    """Create or overwrite a styled Word document (.docx) — v1.7「Office 体系 2.0」字体纪律版.

    Every built-in style (Normal / Title / Heading 1-3) is injected with an EXPLICIT font (token
    body/heading family, all three rFonts links incl. **w:eastAsia** so 中文 no longer falls back to
    宋体), a size ladder (Normal 11 / H3 14 / H2 16 / H1 20 / Title 28 pt), the token heading colour,
    行距 1.4 and加大 H1 前段距. Headings get a coloured 2.25pt hairline underline. So no run is 裸奔 and
    old「字体都不对、中文大小不一」is fixed structurally.

    content markdown-lite (向后兼容 — old plain-text/heading calls work unchanged and auto-inherit the
    new styles):
        '# ' / '## ' / '### '  -> heading levels 1/2/3 (with the accent hairline)
        '- '                   -> bullet point
        '1. ' (2./3.…)         -> numbered point
        blank line             -> spacer

    v1.7 additions:
      * style: 'business'「青花商务」(default) | 'minimal'「墨白极简」| 'vibrant'「活力现代」. Unknown -> business.
      * cover: optional {title, subtitle?, date?, author?} — a full-width 深底满铺 title-block (white 大字
        + 鎏金/强调金 line + subtitle + author·date) on its own page, then a page break before the body.
      * page_numbers: True adds a centred 「第 X 页」 footer via a live PAGE field.
      * inline table 段落: a line 'TABLE: h1 | h2 | h3' begins a token-styled table (表头主色白字 + 斑马纹
        + 细边框); each following '| a | b | c' line is a row; a blank line ends the table. Non-table
        content is unaffected (向后兼容).

    Args:
        path: Output file path (must end with .docx).
        content: Body text in markdown-lite (see above).
        title: Optional document title — added as a Title-styled heading at the top of the body.
        style: Design style name (see above).
        cover: Optional cover-page spec dict (see above).
        page_numbers: Add a 「第 X 页」 page-number footer (default off).
        allow_protected: Override the protected-system-root guard on the destination (default off).

    Returns:
        dict with 'success', 'path', 'output_path', 'style'. On failure {'error': ...}.
    """
    guard = _protected_write_guard(path, allow_protected)
    if guard:
        return guard
    from docx import Document

    tokens = style_tokens.get_style(style)
    resolved_style = style if style in style_tokens.STYLES else style_tokens.DEFAULT_STYLE

    try:
        doc = Document()

        # 字体纪律 backbone: inject explicit fonts/sizes/colours into the built-in styles FIRST,
        # so every paragraph added below inherits a fully specified font (no 裸奔 run).
        _apply_word_styles(doc, tokens)

        if cover and isinstance(cover, dict):
            _add_cover(doc, cover, tokens)

        if title:
            doc.add_heading(str(title), level=0)  # Title style (level 0)

        rule_color = tokens["word_rule_color"]
        # v1.7.1: minimal 标题不用底线而用左侧青色细竖线 (word_heading_rule='left'); 其它风格照旧 bottom.
        rule_side = "left" if tokens.get("word_heading_rule", "bottom") == "left" else "bottom"

        # Table accumulation state: when we hit 'TABLE:' we buffer rows until a blank line.
        pending_table = None  # dict{headers, rows} or None

        def _flush_table():
            nonlocal pending_table
            if pending_table and pending_table["headers"]:
                _add_content_table(doc, pending_table["headers"], pending_table["rows"], tokens)
            pending_table = None

        for line in content.split("\n"):
            stripped = line.strip()

            # --- table block handling ---
            if pending_table is not None:
                if not stripped:
                    _flush_table()
                    doc.add_paragraph("")
                    continue
                if stripped.startswith("|"):
                    cells = [c.strip() for c in stripped.strip("|").split("|")]
                    pending_table["rows"].append(cells)
                    continue
                # a non-row, non-blank line ends the table then falls through to normal handling.
                _flush_table()

            if not stripped:
                doc.add_paragraph("")
            elif stripped.upper().startswith("TABLE:"):
                headers = [c.strip() for c in stripped[6:].split("|") if c.strip()]
                pending_table = {"headers": headers, "rows": []}
            elif stripped.startswith("### "):
                p = doc.add_heading(stripped[4:], level=3)
                _add_heading_rule(p, rule_color, side=rule_side)
            elif stripped.startswith("## "):
                p = doc.add_heading(stripped[3:], level=2)
                _add_heading_rule(p, rule_color, side=rule_side)
            elif stripped.startswith("# "):
                p = doc.add_heading(stripped[2:], level=1)
                _add_heading_rule(p, rule_color, side=rule_side)
            elif stripped.startswith("- "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            elif stripped.startswith("1. ") or stripped.startswith("2. ") or stripped.startswith("3. "):
                doc.add_paragraph(stripped[3:], style="List Number")
            else:
                doc.add_paragraph(stripped)

        _flush_table()  # close a table that ran to EOF

        if page_numbers:
            _add_page_number_footer(doc, tokens)

        os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
        doc.save(path)
        # v1.5.1: echo output_path (== path) so the workbench 产物收割 (ARTIFACT_OUTPUT_PATH_KEYS)
        # picks this file up directly. 老字段 path 保留(字段只增,不破坏现有契约)。
        return {"success": True, "path": os.path.abspath(path),
                "output_path": os.path.abspath(path), "style": resolved_style}
    except Exception as e:
        return {"error": str(e)}


def _looks_numeric(value) -> bool:
    """True if `value` (possibly a numeric string) should be treated as a number for formatting."""
    if isinstance(value, bool):
        return False
    if isinstance(value, (int, float)):
        return True
    if isinstance(value, str):
        s = value.strip().replace(",", "").replace("%", "").replace("¥", "").replace("$", "").strip()
        if s in ("", "-", "."):
            return False
        try:
            float(s)
            return True
        except ValueError:
            return False
    return False


def _numeric_value(value):
    """Coerce a numeric-looking string to int/float so Excel stores a real number (not text)."""
    if isinstance(value, (int, float)) and not isinstance(value, bool):
        return value
    if isinstance(value, str):
        s = value.strip().replace(",", "").replace("%", "").replace("¥", "").replace("$", "").strip()
        try:
            f = float(s)
            return int(f) if f == int(f) else f
        except ValueError:
            return value
    return value


def _column_number_format(header, sample_values) -> str | None:
    """Heuristic Excel number format for a column, from its header text + a sample of its values.

    * 含 '%' 表头 OR values 含 '%'                       -> '0.0%' (values are 0-1 fractions? no —
                                                            values are like '12%' text; we store the
                                                            fraction and format as percent).
      Actually we DON'T rescale — if the source says '12%' we store 0.12 and format 0%. See below.
    * 含 '$'/'¥' 表头, 或 '金额/收入/成本/价格/费用' 关键词  -> '¥#,##0.00' 货币.
    * 纯数字列且任一值 > 999                              -> '#,##0' 千分位.
    * 否则 None (不设格式).
    """
    h = str(header or "")
    money_kw = ("金额", "收入", "成本", "价格", "费用", "支出", "营收", "利润", "销售额", "总额")
    if "%" in h or any(isinstance(v, str) and "%" in v for v in sample_values):
        return "0.0%"
    if "$" in h or "¥" in h or any(k in h for k in money_kw):
        return "¥#,##0.00"
    # plain numeric column with a big value -> thousands separator
    numeric = [v for v in sample_values if _looks_numeric(v)]
    if numeric and len(numeric) >= max(1, len(sample_values) // 2):
        try:
            if any(abs(float(_numeric_value(v))) > 999 for v in numeric):
                return "#,##0"
        except (ValueError, TypeError):
            pass
    return None


@mcp.tool()
def write_excel(
    path: str,
    data: list[list[str]],
    sheet_name: str = "Sheet1",
    headers: list[str] | None = None,
    style: str = "business",
    allow_protected: bool = False,
) -> dict:
    """Create or overwrite a styled Excel file (.xlsx) — v1.7「Office 体系 2.0」落笔即样式版.

    Unlike the old bare writer (which left ZERO font declarations → mixed Calibri/宋体), every cell is
    written with the token body font (微软雅黑, 11pt); the header row (if given) is bold on that font;
    numeric-looking columns get a number format by heuristic (千分位 for values >999, 百分比 for '%'
    columns, 货币 for '$'/'¥'/金额/收入/成本… headers) and numeric-looking strings are coerced to real
    numbers so Excel can compute on them; column widths auto-fit (CJK-aware) on first write. So the
    sheet 「不跑 beautify 也不难看」. excel_beautify remains the full makeover (frozen header, zebra,
    borders, auto-filter) — this is the 落笔即样式 baseline.

    Args:
        path: Output file path (must end with .xlsx).
        data: 2D list of cell values (list of rows, each row is a list of values).
        sheet_name: Name of the worksheet.
        headers: Optional list of header row values (rendered bold on the token font).
        style: Design style — 'business' (default) | 'minimal' | 'vibrant'. Unknown -> business.
        allow_protected: Override the protected-system-root guard on the destination (default off).

    Returns:
        dict with 'success', 'path', 'output_path', 'rows', 'style'. On failure {'error': ...}.
    """
    guard = _protected_write_guard(path, allow_protected)
    if guard:
        return guard
    from openpyxl import Workbook
    from openpyxl.styles import Font, Border, Side
    from openpyxl.utils import get_column_letter

    tokens = style_tokens.get_style(style)
    resolved_style = style if style in style_tokens.STYLES else style_tokens.DEFAULT_STYLE
    body_font_name = tokens["body_font"]
    text_argb = style_tokens.argb(tokens["text_color"])

    try:
        wb = Workbook()
        ws = wb.active
        ws.title = sheet_name

        base_font = Font(name=body_font_name, size=11, color=text_argb)
        header_font = Font(name=body_font_name, size=11, bold=True, color=text_argb)
        # v1.7.1 minimal「墨白极简」落笔即样式: 表头底部 2pt 青色下边框 (excel_header='underline') —
        # 与 excel_beautify 的 minimal 表头语言一致, 让「不跑 beautify」的表也有极简装帧感.
        header_border = None
        if tokens.get("excel_header") == "underline":
            header_border = Border(bottom=Side(style="medium",
                                               color=style_tokens.argb(tokens["accent"])))

        start_row = 1
        n_cols = 0
        if headers:
            n_cols = len(headers)
            for col, header in enumerate(headers, 1):
                c = ws.cell(row=1, column=col, value=header)
                c.font = header_font
                if header_border is not None:
                    c.border = header_border
            start_row = 2

        # write data (coercing numeric-looking strings to real numbers) with the token body font.
        for row_idx, row_data in enumerate(data, start_row):
            n_cols = max(n_cols, len(row_data))
            for col_idx, value in enumerate(row_data, 1):
                c = ws.cell(row=row_idx, column=col_idx, value=_numeric_value(value))
                c.font = base_font

        # per-column number formats (heuristic) + content-fit widths (CJK-aware).
        for col in range(1, n_cols + 1):
            header_text = headers[col - 1] if headers and col - 1 < len(headers) else ""
            sample = [row[col - 1] for row in data if col - 1 < len(row)][:50]
            fmt = _column_number_format(header_text, sample)
            longest = 0
            # header width
            if headers and col - 1 < len(headers):
                longest = sum(2 if ord(ch) > 0x2E7F else 1 for ch in str(headers[col - 1]))
            for r in range(start_row, start_row + len(data)):
                cell = ws.cell(row=r, column=col)
                if fmt is not None and _looks_numeric(cell.value):
                    # percentage: values like '12%' were coerced to 12 (not 0.12); if the source was
                    # a fraction string we keep it. We store the raw number and format; for '12%' text
                    # coerced to 12, format 0.0% would show 1200% — so guard: only apply % when the
                    # coerced value is already a 0-1 fraction. Otherwise fall back to plain number.
                    if fmt == "0.0%" and isinstance(cell.value, (int, float)) and abs(cell.value) > 1:
                        cell.number_format = "0.0\"%\""  # show the number with a literal % suffix
                    else:
                        cell.number_format = fmt
                v = cell.value
                if v is not None:
                    disp = sum(2 if ord(ch) > 0x2E7F else 1 for ch in str(v))
                    # 把关直修(v1.7 审美关真机撞出):数字格式会加宽显示——"1250" 套上 ¥#,##0.00 变成
                    # "¥1,250.00",列宽若按【原始值】量,渲染出来就是 ######。这里按格式化后的近似宽度
                    # (千分位分隔符 + 小数位 + 货币/百分号 + 负号)取更大者。近似即可,+2 padding 兜底。
                    if fmt is not None and isinstance(cell.value, (int, float)):
                        f = str(cell.number_format or "")
                        try:
                            dec = 2 if "0.00" in f else (1 if "0.0" in f else 0)
                            body = f"{abs(float(cell.value)):,.{dec}f}"
                            symbol = 1 if any(s in f for s in ("¥", "$", "％", "%")) else 0
                            sign = 1 if float(cell.value) < 0 else 0
                            disp = max(disp, len(body) + symbol + sign)
                        except Exception:
                            pass
                    longest = max(longest, disp)
            ws.column_dimensions[get_column_letter(col)].width = min(max(longest + 2, 8), 50)

        os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
        wb.save(path)
        # v1.5.1: 补 output_path(== path), 见 write_document。
        return {"success": True, "path": os.path.abspath(path),
                "output_path": os.path.abspath(path), "rows": len(data), "style": resolved_style}
    except Exception as e:
        return {"error": str(e)}


# --- PDF export (write_pdf) ---------------------------------------------------
# reportlab's default fonts (Helvetica et al.) carry no CJK glyphs, so Chinese text
# renders as tofu boxes unless we register a real CJK font. We resolve a font ONCE
# (module-level cache) via a fixed preference chain and reuse it for every export.
_PDF_FONT_CACHE: dict | None = None


def _resolve_cjk_font():
    """Register and cache a Chinese-capable font for reportlab. Returns a dict:

        {"name": <registered font name>, "warning": <optional str>}

    Preference chain (first that works wins):
      1. C:\\Windows\\Fonts\\msyh.ttc   (Microsoft YaHei, TTFont subfontIndex=0)
      2. C:\\Windows\\Fonts\\simsun.ttc (SimSun,          TTFont subfontIndex=0)
      3. reportlab built-in UnicodeCIDFont('STSong-Light') — zero external files
         (the CID font is resolved by the PDF *reader*)
      4. Helvetica (Latin only, extreme fallback) + a warning that Chinese may not show
    """
    global _PDF_FONT_CACHE
    if _PDF_FONT_CACHE is not None:
        return _PDF_FONT_CACHE

    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont

    # 1 & 2: TrueType collections shipped with Windows.
    for font_name, ttc_path in (
        ("MSYaHei", r"C:\Windows\Fonts\msyh.ttc"),
        ("SimSun", r"C:\Windows\Fonts\simsun.ttc"),
    ):
        if os.path.exists(ttc_path):
            try:
                pdfmetrics.registerFont(TTFont(font_name, ttc_path, subfontIndex=0))
                _PDF_FONT_CACHE = {"name": font_name}
                return _PDF_FONT_CACHE
            except Exception:
                # Corrupt/unsupported collection -> fall through to the next candidate.
                pass

    # 3: reportlab's built-in Adobe CID font — no external file needed.
    try:
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        _PDF_FONT_CACHE = {"name": "STSong-Light"}
        return _PDF_FONT_CACHE
    except Exception:
        pass

    # 4: extreme fallback — Latin only. Chinese will not render.
    _PDF_FONT_CACHE = {
        "name": "Helvetica",
        "warning": "未找到中文字体，中文可能无法显示",
    }
    return _PDF_FONT_CACHE


@mcp.tool()
def write_pdf(
    path: str,
    content: str,
    title: str | None = None,
    table_headers: list[str] | None = None,
    table_data: list[list[str]] | None = None,
    page_size: str = "A4",
    allow_protected: bool = False,
) -> dict:
    """Create or overwrite a PDF (.pdf) from markdown-lite text, with full Chinese support.

    content uses the same markdown-lite syntax as write_document:
        '# ' / '## ' / '### '  -> heading levels 1/2/3
        '- '                   -> bullet point
        '1. ' (or '2.'/'3.')   -> numbered point
        blank line             -> vertical spacing

    If table_data is given, a table is rendered after the body (table_headers optional
    as the header row). Every cell is str()-coerced.

    A Chinese-capable font is auto-registered (Microsoft YaHei -> SimSun -> reportlab's
    built-in STSong-Light CID font -> Helvetica as a last resort). Paragraphs use
    wordWrap='CJK' so Chinese lines break correctly.

    Args:
        path: Output file path (must end with .pdf).
        content: Body text in markdown-lite (see above).
        title: Optional document title (rendered as the top heading).
        table_headers: Optional header row for the trailing table.
        table_data: Optional 2D list of rows for the trailing table.
        page_size: 'A4' or 'letter' (anything else falls back to A4).

    Returns:
        dict with 'success', 'path' (abs), 'pages', 'font'. On the Helvetica fallback
        it also carries 'warning'. Missing reportlab -> {'error': install guidance}.
    """
    if not str(path).lower().endswith(".pdf"):
        return {"error": "path must end with .pdf"}
    guard = _protected_write_guard(path, allow_protected)
    if guard:
        return guard

    # Lazy, guarded import — reportlab is an OPTIONAL offline dependency. Absent -> degrade.
    try:
        from reportlab.lib.pagesizes import A4, letter
        from reportlab.lib.units import mm
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import (
            SimpleDocTemplate,
            Paragraph,
            Spacer,
            Table,
            TableStyle,
        )
    except Exception:
        return {"error": "PDF 导出需要 reportlab。离线包已含，可运行 installer 重装；或 pip install reportlab"}

    try:
        font_info = _resolve_cjk_font()
        font_name = font_info["name"]

        # Paragraph() parses mini-HTML, so raw '&'/'<'/'>' in USER text (e.g. "R&D",
        # "<url>") would raise or misparse. Escape every user-supplied string; the only
        # markup we ever emit (<b> header cells) is our own, wrapped OUTSIDE the escape.
        def _esc(s) -> str:
            return str(s).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

        # page size: only 'A4' | 'letter'; anything else -> A4.
        psize = letter if str(page_size).lower() == "letter" else A4

        # Clone the sample stylesheet but force every style onto our CJK font, with
        # wordWrap='CJK' so Chinese wraps mid-run (no whitespace to break on).
        base = getSampleStyleSheet()

        def _cjk_style(src_name: str, **overrides) -> ParagraphStyle:
            src = base[src_name]
            return ParagraphStyle(
                f"CJK-{src_name}",
                parent=src,
                fontName=font_name,
                wordWrap="CJK",
                **overrides,
            )

        style_body = _cjk_style("BodyText")
        style_title = _cjk_style("Title")
        style_h1 = _cjk_style("Heading1")
        style_h2 = _cjk_style("Heading2")
        style_h3 = _cjk_style("Heading3")
        style_bullet = _cjk_style("BodyText", leftIndent=18)

        story = []

        if title:
            story.append(Paragraph(_esc(title), style_title))
            story.append(Spacer(1, 6 * mm))

        for raw in str(content).split("\n"):
            stripped = raw.strip()
            if not stripped:
                story.append(Spacer(1, 4 * mm))
            elif stripped.startswith("### "):
                story.append(Paragraph(_esc(stripped[4:]), style_h3))
            elif stripped.startswith("## "):
                story.append(Paragraph(_esc(stripped[3:]), style_h2))
            elif stripped.startswith("# "):
                story.append(Paragraph(_esc(stripped[2:]), style_h1))
            elif stripped.startswith("- "):
                story.append(Paragraph("• " + _esc(stripped[2:]), style_bullet))
            elif stripped.startswith(("1. ", "2. ", "3. ", "4. ", "5. ",
                                      "6. ", "7. ", "8. ", "9. ")):
                story.append(Paragraph(stripped[0:2] + " " + _esc(stripped[3:]), style_bullet))
            else:
                story.append(Paragraph(_esc(stripped), style_body))

        # Optional trailing table. Every cell wrapped in a CJK Paragraph so wide/Chinese
        # cells wrap inside the column instead of overflowing the page.
        if table_data:
            rows = []
            if table_headers:
                rows.append([Paragraph("<b>" + _esc(h) + "</b>", style_body) for h in table_headers])
            for row in table_data:
                rows.append([Paragraph(_esc(c), style_body) for c in row])
            if rows:
                story.append(Spacer(1, 4 * mm))
                tbl = Table(rows, repeatRows=1 if table_headers else 0)
                tbl.setStyle(
                    TableStyle(
                        [
                            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                            ("VALIGN", (0, 0), (-1, -1), "TOP"),
                            ("BACKGROUND", (0, 0), (-1, 0),
                             colors.whitesmoke if table_headers else colors.white),
                            ("FONTNAME", (0, 0), (-1, -1), font_name),
                            ("LEFTPADDING", (0, 0), (-1, -1), 4),
                            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                            ("TOPPADDING", (0, 0), (-1, -1), 3),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                        ]
                    )
                )
                story.append(tbl)

        os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)

        # Count pages via an onPage callback (robust across reportlab builds).
        _page_counter = {"n": 0}

        def _count_page(canvas, doc):
            _page_counter["n"] += 1

        doc = SimpleDocTemplate(path, pagesize=psize)
        doc.build(story, onFirstPage=_count_page, onLaterPages=_count_page)
        pages = _page_counter["n"] or getattr(doc, "page", 1) or 1

        # v1.5.1: 补 output_path(== path), 见 write_document。
        out = {"success": True, "path": os.path.abspath(path), "output_path": os.path.abspath(path), "pages": pages, "font": font_name}
        if font_info.get("warning"):
            out["warning"] = font_info["warning"]
        return out
    except Exception as e:
        return {"error": str(e)}
